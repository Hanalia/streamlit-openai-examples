import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup


def markdown_to_docx(input_file, output_file):
    # Read the Markdown content
    with open(input_file, "r", encoding="utf-8") as f:
        markdown_content = f.read()

    # Convert Markdown to HTML
    html_content = markdown.markdown(markdown_content)

    # Parse the HTML content
    soup = BeautifulSoup(html_content, "html.parser")

    # Create a new Word document
    doc = Document()

    # Define custom styles
    styles = doc.styles

    def create_heading_style(name, size, color=RGBColor(0, 0, 0), bold=True):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = "Malgun Gothic"  # A common Korean font
        font.size = Pt(size)
        font.color.rgb = color
        font.bold = bold

    create_heading_style("Custom Heading 1", 18, RGBColor(0, 0, 255))  # Blue
    create_heading_style("Custom Heading 2", 16, RGBColor(0, 0, 0))  # Green
    create_heading_style("Custom Heading 3", 12, RGBColor(0, 0, 0))  # Dark Red

    # Function to add content to the document
    def add_element(element):
        if element.name == "p":
            paragraph = doc.add_paragraph(element.text)
            paragraph.style = styles["Normal"]
        elif element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            paragraph = doc.add_paragraph(element.text)
            if level <= 3:
                paragraph.style = styles[f"Custom Heading {level}"]
            else:
                paragraph.style = styles[f"Heading {level}"]
        elif element.name == "ul":
            for li in element.find_all("li"):
                paragraph = doc.add_paragraph(li.text, style="List Bullet")
        elif element.name == "ol":
            for li in element.find_all("li"):
                paragraph = doc.add_paragraph(li.text, style="List Number")
        # Add more conditions for other elements as needed

    # Process each element in the HTML
    for element in soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol"]):
        add_element(element)

    # Save the document
    doc.save(output_file)


# Usage example
markdown_to_docx("input.md", "output.docx")

from io import BytesIO
import os
import streamlit as st
from openai import OpenAI
from dotenv import load_dotenv
import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup


load_dotenv()


## 다운로드 버튼으로 전체 페이지가 reload 되는 것을 방지
@st.fragment
def show_download_button(bytes: BytesIO):
    st.download_button(
        label="보고서 워드파일 다운로드",
        data=bytes,
        file_name="Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def markdown_to_docx(markdown_content: str) -> BytesIO:

    # Convert Markdown to HTML
    html_content = markdown.markdown(markdown_content)

    # Parse the HTML content
    soup = BeautifulSoup(html_content, "html.parser")

    # Create a new Word document
    doc = Document()

    # Define custom styles
    styles = doc.styles

    def create_heading_style(name, size, color=RGBColor(0, 0, 0), bold=True):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = "Batang"  # A common Korean font "Malgun Gothic"
        font.size = Pt(size)
        font.color.rgb = color
        font.bold = bold

    create_heading_style("Custom Heading 1", 18, RGBColor(0, 0, 255))  # Blue
    create_heading_style("Custom Heading 2", 16, RGBColor(0, 0, 0))  # Green
    create_heading_style("Custom Heading 3", 12, RGBColor(0, 0, 0))  # Dark Red

    # Function to add content to the document
    def add_element(element):
        if element.name == "p":
            paragraph = doc.add_paragraph(element.text)
            paragraph.style = styles["Normal"]
        elif element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            paragraph = doc.add_paragraph(element.text)
            if level <= 3:
                paragraph.style = styles[f"Custom Heading {level}"]
            else:
                paragraph.style = styles[f"Heading {level}"]
        elif element.name == "ul":
            for li in element.find_all("li"):
                paragraph = doc.add_paragraph(li.text, style="List Bullet")
        elif element.name == "ol":
            for li in element.find_all("li"):
                paragraph = doc.add_paragraph(li.text, style="List Number")
        # Add more conditions for other elements as needed

    # Process each element in the HTML
    for element in soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6", "ul", "ol"]):
        add_element(element)
    # Convert the document to a byte stream
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io


def main():
    st.set_page_config(layout="wide")
    st.title("보고서 작성")

    with st.sidebar:
        openai_api_key = st.text_input(
            "OpenAI API Key", key="summarizer_api_key", type="password"
        )

        st.markdown(
            "[OpenAI API key 받기](https://platform.openai.com/account/api-keys)"
        )

    ## 개반전용 모드
    if os.environ["DEV_MODE"] == "TRUE":
        openai_api_key = os.getenv("OPENAI_API_KEY")

    system_message = """
    너는 보고서 작성 전문가야.
    - 마크다운으로 보고서 작성을 체계적으로 해줘
    - 3개의 heading2 그리고 각각 2개의 heading3로 구성해줘
    - heading 2의 내용은 300자 이상으로 해줘
    - 목차는 제외해줘
    - 보고서 내용만 응답 결과로 표현해줘
    """

    if "messages" not in st.session_state:
        st.session_state.messages = [{"role": "system", "content": system_message}]

    default_user_input = """생성형 AI가 세상을 어떻게 바꿀 수 있을까?"""

    user_input = st.text_area(
        "작성할 보고서의 주요 주제, 내용을 입력하세요:",
        value=default_user_input,
        height=70,
    )

    if st.button("보고서 초안 작성하기"):
        if not openai_api_key:
            st.info("계속하려면 OpenAI API 키를 추가하세요.")
            st.stop()

        if not user_input.strip():
            st.warning("요약할 텍스트를 입력해주세요.")
            st.stop()

        st.session_state["client"] = OpenAI(api_key=openai_api_key)

        prompt = f"다음 내용을 바탕으로 보고서를 작성해 줘:\n\n{user_input}"
        st.session_state.messages.append({"role": "user", "content": prompt})

        with st.spinner("작성 중..."):
            stream = st.session_state["client"].chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": m["role"], "content": m["content"]}
                    for m in st.session_state.messages
                ],
                stream=True,
            )
            response = st.write_stream(stream)
            # Convert Markdown response to a DOCX file in-memory
            docx_content_in_bytes = markdown_to_docx(response)

            show_download_button(docx_content_in_bytes)
            # st.download_button(
            #     label="보고서 워드파일 다운로드",
            #     data=docx_file,
            #     file_name="Report.docx",
            #     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            # )


if __name__ == "__main__":
    main()
